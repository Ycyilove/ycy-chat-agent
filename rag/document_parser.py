"""
文档解析模块
支持解析PDF、TXT、DOCX格式文件，自动去除空白、乱码、冗余换行
"""
import os
import re
from typing import List, Optional
from io import BytesIO


class DocumentParser:
    """文档解析器基类"""

    def parse(self, file_content: bytes, filename: str) -> str:
        """解析文件内容"""
        raise NotImplementedError

    def _clean_text(self, text: str) -> str:
        """清理文本：去除多余空白和乱码"""
        if not text:
            return ""
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


class PDFParser(DocumentParser):
    """PDF文档解析器"""

    def __init__(self):
        try:
            from pypdf import PdfReader
            self.PdfReader = PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                self.PdfReader = PdfReader
            except ImportError:
                raise ImportError("请安装 pypdf 或 PyPDF2: pip install pypdf")

    def parse(self, file_content: bytes, filename: str) -> str:
        """解析PDF文件"""
        pdf_file = BytesIO(file_content)
        reader = self.PdfReader(pdf_file)
        text_parts = []

        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text()
            if text:
                text_parts.append(text)

        full_text = "\n".join(text_parts)
        return self._clean_text(full_text)


class TXTParser(DocumentParser):
    """TXT文档解析器"""

    def parse(self, file_content: bytes, filename: str) -> str:
        """解析TXT文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5']

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except (UnicodeDecodeError, AttributeError):
                continue
        else:
            text = file_content.decode('utf-8', errors='ignore')

        return self._clean_text(text)


class DOCXParser(DocumentParser):
    """DOCX文档解析器"""

    def __init__(self):
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

    def parse(self, file_content: bytes, filename: str) -> str:
        """解析DOCX文件"""
        docx_file = BytesIO(file_content)
        doc = self.Document(docx_file)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        return self._clean_text(full_text)


class DocumentParserFactory:
    """文档解析器工厂"""

    _parsers = {
        '.pdf': PDFParser,
        '.txt': TXTParser,
        '.docx': DOCXParser,
    }

    @classmethod
    def get_parser(cls, file_extension: str) -> Optional[DocumentParser]:
        """获取对应的解析器"""
        parser_class = cls._parsers.get(file_extension.lower())
        if parser_class:
            return parser_class()
        return None

    @classmethod
    def parse_file(cls, file_content: bytes, filename: str) -> str:
        """解析文件"""
        ext = os.path.splitext(filename)[1]
        parser = cls.get_parser(ext)

        if parser is None:
            raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {', '.join(cls._parsers.keys())}")

        return parser.parse(file_content, filename)

    @classmethod
    def supported_formats(cls) -> List[str]:
        """获取支持的文件格式"""
        return list(cls._parsers.keys())
